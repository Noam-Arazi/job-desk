"""The tailoring agent, and one rejection test per rule in the change contract.

The suite is the deliverable, not the proof that the deliverable works. Every
guarantee in spec/change-contract.yaml is a claim about what the system will
refuse, and a refusal that has never been demonstrated is an assumption. So
each `forbidden` id below gets a changeset that breaks it and an assertion that
the run fails naming that id.

Everything is built here, in tmp_path. The six approved bases and the
experience inventory live outside the repo because they carry family and client
detail, and nothing from them is copied into a test, a fixture or a cassette.
The synthetic base is assembled from the contract's own anchor wording, which
is public, so the anchor rules are exercised against the same strings the
production code will look for.
"""

from __future__ import annotations

import copy
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pytest
from docx import Document

import desk.tailor.tailor as tl
from desk.analyst.types import Analysis, Family, Fit, Requirement
from desk.evals.guardrails import _REACH_PARAMS
from desk.llm.base import LLMResponse
from desk.registry import registry
from desk.tailor import bases, render
from desk.tailor import changeset as cs
from desk.tailor import contract as ct
from desk.trace import Usage

CONTRACT = ct.load_contract()
ANCHORS = {a["id"]: a for a in CONTRACT["anchors"]}

SUMMARY_TEXT = "A locked summary that says what was actually managed."
BULLET_DASHBOARDS = "Built dashboards for the weekly review"
BULLET_PROTOTYPE = "Built a retrieval prototype for one client"
SKILL_DATA = "Data: sql, python"
SKILL_TOOLS = "Tools: monday, make"
SKILL_CLOUD = "Cloud: vertex ai, firebase"


# --- a base, built from nothing ------------------------------------------


def _bold(document: Any, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(text).bold = True


def _skills_line(document: Any, text: str) -> None:
    """A skills line the way the approved bases carry it: two runs, not one.

    The bold category and the plain items are separate runs, which is the whole
    reason `render._set_text` localises an edit instead of rewriting the
    paragraph. Built as a single run, this line would survive total flattening
    unchanged and every test about run formatting would hold vacuously.
    """
    head, _, items = text.partition(": ")
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{head}: ").bold = True
    paragraph.add_run(items)


def make_base_file(directory: Path, name: str = "Noam_Arazi_CV_AI_Builder_base.docx") -> Path:
    """A synthetic .docx with the same shape as an approved base.

    Shape only: an identity block, a locked summary, three skills lines, two
    employers with ordered bullets, and a structural education section. The two
    employers are the ones the contract names anchors on.
    """
    document = Document()
    _bold(document, "Test Person")
    document.add_paragraph("Haifa | +972-00-000-0000 | test@example.com")
    _bold(document, "Summary")
    document.add_paragraph(SUMMARY_TEXT)
    _bold(document, "Technical Skills")
    for line in (SKILL_DATA, SKILL_TOOLS, SKILL_CLOUD):
        _skills_line(document, line)
    _bold(document, "Experience")
    _bold(document, "Growth Directorate  |  2024 - 2026")
    document.add_paragraph("Information systems")
    document.add_paragraph(ANCHORS["growth_data_layer"]["en"], style="List Bullet")
    document.add_paragraph(ANCHORS["growth_analyst_team"]["en"], style="List Bullet")
    document.add_paragraph(BULLET_DASHBOARDS, style="List Bullet")
    _bold(document, "Fischer Technologies  |  2025 - 2026")
    document.add_paragraph("AI solutions")
    document.add_paragraph(ANCHORS["fischer_ai_solutions"]["en"], style="List Bullet")
    document.add_paragraph(BULLET_PROTOTYPE, style="List Bullet")
    _bold(document, "Education")
    document.add_paragraph("University of Haifa  |  2022 - 2026")

    directory.mkdir(parents=True, exist_ok=True)
    path = directory / name
    document.save(str(path))
    return path


@pytest.fixture
def bases_dir(tmp_path) -> Path:
    directory = tmp_path / "bases"
    make_base_file(directory)
    return directory


@pytest.fixture
def base(bases_dir) -> bases.Base:
    return bases.load_for("ai_builder", directory=bases_dir)


@pytest.fixture
def inventory(tmp_path) -> Path:
    path = tmp_path / "experience.md"
    path.write_text(
        "Reporting dashboards were rebuilt during the quarterly review.\n"
        "Retrieval work used an open-source vector store.\n",
        encoding="utf-8",
    )
    return path


def change(**overrides: Any) -> cs.Change:
    fields: dict[str, Any] = {
        "op": cs.SWAP_TERMINOLOGY,
        "section": "exp.0.bullet.2",
        "before": BULLET_DASHBOARDS,
        "after": BULLET_DASHBOARDS,
        "source": cs.BASE,
        "source_line": BULLET_DASHBOARDS,
    }
    fields.update(overrides)
    return cs.Change(**fields)


def violations_of(base: bases.Base, *changes: cs.Change) -> set[str]:
    return {v.rule for v in ct.check(base, cs.ChangeSet(changes=changes), contract=CONTRACT)}


# --- selecting and loading a base ----------------------------------------


def test_catalog_reads_family_and_language_off_the_filename(tmp_path):
    directory = tmp_path / "six"
    for name in (
        "Noam_Arazi_CV_AI_Builder_base.docx",
        "Noam_Arazi_CV_Data_Analyst_base.docx",
        "Noam_Arazi_CV_Product_Project_base.docx",
        "נעם_ארזי_קורות_חיים_אסטרטגיה_בסיס.docx",
        "נעם_ארזי_קורות_חיים_דאטה_בסיס.docx",
        "נעם_ארזי_קורות_חיים_פרודקט_בסיס.docx",
    ):
        make_base_file(directory, name)

    found = bases.catalog(directory)
    assert len(found) == 6, "six bases, not eight"
    assert {(f.family, f.language) for f in found} == {
        ("ai_builder", "en"),
        ("data_analyst", "en"),
        ("data_analyst", "he"),
        ("product_project", "en"),
        ("product_project", "he"),
        ("strategy_public", "he"),
    }


def test_a_one_language_family_settles_the_language_question(tmp_path):
    """ai_builder is English-only and strategy_public is Hebrew-only — session 2."""
    directory = tmp_path / "six"
    make_base_file(directory, "Noam_Arazi_CV_AI_Builder_base.docx")
    make_base_file(directory, "נעם_ארזי_קורות_חיים_אסטרטגיה_בסיס.docx")

    assert bases.select("ai_builder", directory=directory, language="he").language == "en"
    assert bases.select("strategy_public", directory=directory, language="en").language == "he"


def test_language_defaults_to_hebrew_and_switches_on_an_english_posting():
    assert bases.choose_language(("en", "he")) == "he"
    assert bases.choose_language(("en", "he"), posting_is_english=True) == "en"
    assert bases.choose_language(("en", "he"), international_product_company=True) == "en"
    assert bases.choose_language(("en", "he"), requested="en") == "en"


def test_an_unknown_family_is_an_error_not_a_guess(bases_dir):
    with pytest.raises(bases.BaseNotFound):
        bases.select("microbiologist", directory=bases_dir)


def test_parsing_addresses_every_line_by_where_it_sits(base):
    kinds = {line.address: line.kind for line in base.lines}
    assert kinds["identity.0"] == bases.IDENTITY
    assert kinds["summary.0"] == bases.SUMMARY
    assert kinds["skills.1"] == bases.SKILL
    assert kinds["exp.0.header"] == bases.EMPLOYER
    assert kinds["exp.0.bullet.0"] == bases.BULLET
    assert kinds["structural.0"] == bases.STRUCTURAL
    assert base.summary_text == SUMMARY_TEXT
    assert base.employer_of("exp.1.bullet.0") == "fischer_technologies"
    assert base.employer_of("exp.0.bullet.0") == "growth_directorate"


def test_only_skills_and_bullets_are_tailorable(base):
    tailorable = {line.kind for line in base.lines if line.tailorable}
    assert tailorable == {bases.SKILL, bases.BULLET}


def test_the_base_is_re_read_and_re_hashed_every_run(bases_dir):
    """Noam edits the bases by hand in Word; a cached copy would hide his edit."""
    first = bases.load_for("ai_builder", directory=bases_dir)
    assert first.line("skills.0").text == SKILL_DATA

    edited = Document(str(first.path))  # stand in for Noam editing it in Word
    for paragraph in edited.paragraphs:
        if paragraph.text == SKILL_DATA:
            paragraph.runs[0].text = "Data: sql, python, duckdb"
            for run in paragraph.runs[1:]:
                run.text = ""
    edited.save(str(first.path))

    second = bases.load_for("ai_builder", directory=bases_dir)
    assert second.line("skills.0").text == "Data: sql, python, duckdb"
    assert second.sha256 != first.sha256
    assert second.sha256 == bases.sha256_of(second.path)


def test_every_anchor_is_located_in_the_base(base):
    located = {a.id: a.address for a in ct.locate_anchors(base, CONTRACT)}
    assert located["growth_data_layer"] == "exp.0.bullet.0"
    assert located["growth_analyst_team"] == "exp.0.bullet.1"
    assert located["fischer_ai_solutions"] == "exp.1.bullet.0"


def make_confusable_base(directory: Path) -> Path:
    """A base where another employer's bullet is the best match for an anchor.

    Noam reworded the Growth bullet by hand, and the Fischer section happens to
    describe the same kind of work. Nothing here is malformed; the file is
    exactly the sort of thing hand-editing produces.
    """
    document = Document()
    _bold(document, "Test Person")
    document.add_paragraph("Haifa | test@example.com")
    _bold(document, "Summary")
    document.add_paragraph(SUMMARY_TEXT)
    _bold(document, "Technical Skills")
    _skills_line(document, SKILL_DATA)
    _bold(document, "Experience")
    _bold(document, "Growth Directorate  |  2024 - 2026")
    document.add_paragraph("Information systems")
    document.add_paragraph("Ran the weekly reporting flow", style="List Bullet")
    _bold(document, "Fischer Technologies  |  2025 - 2026")
    document.add_paragraph("AI solutions")
    document.add_paragraph(
        ANCHORS["growth_data_layer"]["en"] + " for one client", style="List Bullet"
    )
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / "Noam_Arazi_CV_AI_Builder_base.docx"
    document.save(str(path))
    return path


def test_an_anchor_binds_only_to_a_line_of_its_own_employer(tmp_path):
    """`employer:` is part of the anchor, and the locator ignored it.

    An anchor bound to another employer's bullet reports itself protected while
    the real line it names is free to be dropped — the guarantee inverted.
    """
    base = bases.load(make_confusable_base(tmp_path / "bases"))
    growth = {a.id: a for a in ct.locate_anchors(base, CONTRACT)}["growth_data_layer"]

    assert growth.employer == "growth_directorate"
    assert base.employer_of("exp.1.bullet.0") == "fischer_technologies"
    assert growth.address != "exp.1.bullet.0", "the anchor bound to another employer's line"
    assert not growth.located, "this base no longer phrases the claim; nothing here is protected"


# --- projecting a changeset ----------------------------------------------


def test_projection_edits_drops_and_reorders_without_touching_a_file(base):
    projected = cs.project(
        base,
        cs.ChangeSet(
            changes=(
                change(section="skills.0", after="Data: SQL, Python"),
                change(op=cs.DROP_SECONDARY_BULLET, section="exp.1.bullet.1", after=""),
                change(
                    op=cs.REORDER_SKILL_CATEGORIES,
                    section="skills",
                    before="skills.0, skills.1, skills.2",
                    after="skills.2, skills.0, skills.1",
                ),
            )
        ),
    )
    assert projected.of("skills.0") == "Data: SQL, Python"
    assert projected.of("exp.1.bullet.1") is None
    order = [a for a in projected.addresses if a.startswith("skills.")]
    assert order == ["skills.2", "skills.0", "skills.1"]
    assert projected.of("summary.0") == SUMMARY_TEXT


# --- the contract: one rejection per forbidden rule -----------------------


def test_every_forbidden_rule_has_a_deterministic_check():
    """A rule in the file with nothing behind it is a contract that is not enforced."""
    missing = set(ct.forbidden_ids(CONTRACT)) - set(ct.CHECKS)
    assert not missing, f"no check implements {sorted(missing)}"


# One changeset per forbidden id that breaks that rule. The named tests
# below carry the nuance — the carve-outs, the parametrised phrase lists — and
# this table carries the guarantee the session owes: a rule added to the YAML
# with nothing that demonstrates its refusal fails here, by id, immediately.
REJECTIONS: dict[str, tuple[dict[str, Any], ...]] = {
    "edit_summary": ({"section": "summary.0", "after": "A per-posting summary."},),
    "add_new_bullet": ({"section": "exp.0.bullet.9", "before": "", "after": "A new bullet"},),
    "introduce_number": ({"after": "Built 12 dashboards for the weekly review"},),
    "exceed_one_page": ({"after": BULLET_DASHBOARDS + (" and much more besides" * 20)},),
    "drop_or_weaken_anchor": (
        {"op": cs.DROP_SECONDARY_BULLET, "section": "exp.0.bullet.0", "after": ""},
    ),
    "claim_without_evidence": ({"after": "Built reporting dashboards", "source": ""},),
    "fischer_adoption_claim": (
        {
            "section": "exp.1.bullet.1",
            "before": BULLET_PROTOTYPE,
            "after": BULLET_PROTOTYPE + ", in production",
        },
    ),
    "attribute_make_to_fischer": (
        {
            "section": "exp.1.bullet.1",
            "before": BULLET_PROTOTYPE,
            "after": BULLET_PROTOTYPE + " with Make",
        },
    ),
    "vertex_on_fischer": (
        {
            "section": "exp.1.bullet.1",
            "before": BULLET_PROTOTYPE,
            "after": BULLET_PROTOTYPE + " on Vertex AI",
        },
    ),
    "write_ai_assisted": ({"after": "AI-assisted dashboards for the weekly review"},),
    "touch_identity_block": ({"section": "identity.1", "before": "", "after": "a new number"},),
}


@pytest.mark.parametrize("rule_id", ct.forbidden_ids(CONTRACT))
def test_every_forbidden_rule_rejects_a_changeset_that_breaks_it(base, rule_id):
    assert rule_id in REJECTIONS, f"{rule_id} has no changeset demonstrating its refusal"
    caught = violations_of(base, *(change(**c) for c in REJECTIONS[rule_id]))
    assert rule_id in caught, f"{rule_id} did not fire; got {sorted(caught)}"


def test_a_clean_changeset_passes(base):
    approved = cs.ChangeSet(
        changes=(
            change(
                op=cs.ADD_TERM_TO_EXISTING_LINE,
                section="skills.0",
                before=SKILL_DATA,
                after=SKILL_DATA + ", dbt",
                source=cs.INVENTORY,
                source_line="Reporting dashboards were rebuilt during the quarterly review.",
            ),
            change(
                section="exp.0.bullet.2",
                before=BULLET_DASHBOARDS,
                after="Built reporting dashboards for the weekly review",
                source=cs.INVENTORY,
                source_line="Reporting dashboards were rebuilt during the quarterly review.",
            ),
            change(
                op=cs.REORDER_SKILL_CATEGORIES,
                section="skills",
                before="skills.0, skills.1, skills.2",
                after="skills.2, skills.0, skills.1",
            ),
        )
    )
    assert ct.check(base, approved, contract=CONTRACT) == ()
    assert ct.enforce(base, approved, contract=CONTRACT).of("summary.0") == SUMMARY_TEXT


def test_edit_summary_is_rejected(base):
    assert "edit_summary" in violations_of(
        base,
        change(section="summary.0", before=SUMMARY_TEXT, after="A summary rewritten per posting."),
    )


def test_edit_summary_catches_a_reorder_that_moves_it(base):
    assert "edit_summary" in violations_of(
        base,
        change(
            op=cs.REORDER_LINES_WITHIN_SECTION,
            section="skills",
            before="summary.0, skills.0",
            after="skills.0, summary.0",
        ),
    )


def test_add_new_bullet_is_rejected(base):
    assert "add_new_bullet" in violations_of(
        base,
        change(section="exp.0.bullet.9", before="", after="A bullet the base never had"),
    )


def test_adding_a_term_may_not_quietly_replace_the_line(base):
    """The `never: creates_a_new_line` clause of add_term_to_existing_line."""
    assert "add_new_bullet" in violations_of(
        base,
        change(
            op=cs.ADD_TERM_TO_EXISTING_LINE,
            section="skills.0",
            before=SKILL_DATA,
            after="Data: airflow",
            source=cs.INVENTORY,
            source_line="airflow",
        ),
    )


def test_introduce_number_is_rejected(base):
    assert "introduce_number" in violations_of(
        base,
        change(
            section="exp.0.bullet.2",
            before=BULLET_DASHBOARDS,
            after="Built 12 dashboards for the weekly review",
        ),
    )


def test_exceed_one_page_is_rejected(base):
    per_line = CONTRACT["length"]["chars_per_line"]
    assert per_line == 108
    assert "exceed_one_page" in violations_of(
        base,
        change(
            section="exp.0.bullet.2",
            before=BULLET_DASHBOARDS,
            after=BULLET_DASHBOARDS + (" and much more besides" * 20),
        ),
    )


def test_dropping_an_anchor_is_rejected(base):
    assert "drop_or_weaken_anchor" in violations_of(
        base,
        change(op=cs.DROP_SECONDARY_BULLET, section="exp.0.bullet.0", after=""),
    )


def test_rewriting_an_anchor_past_recognition_is_rejected(base):
    assert "drop_or_weaken_anchor" in violations_of(
        base,
        change(
            section="exp.0.bullet.0",
            before=ANCHORS["growth_data_layer"]["en"],
            after="Helped out around the office",
        ),
    )


def test_moving_an_anchor_off_its_position_is_rejected(base):
    """`position: first_bullet` is part of the anchor, not decoration."""
    assert "drop_or_weaken_anchor" in violations_of(
        base,
        change(
            op=cs.REORDER_LINES_WITHIN_SECTION,
            section="exp.0",
            before="exp.0.bullet.0, exp.0.bullet.1",
            after="exp.0.bullet.1, exp.0.bullet.0",
        ),
    )


def test_a_change_with_no_source_fails_the_run(base):
    assert "claim_without_evidence" in violations_of(
        base,
        change(section="exp.0.bullet.2", after="Built reporting dashboards", source=""),
    )


def test_a_swap_that_quotes_no_source_line_fails_the_run(base):
    assert "claim_without_evidence" in violations_of(
        base,
        change(
            section="exp.0.bullet.2",
            after="Built reporting dashboards",
            source=cs.INVENTORY,
            source_line="   ",
        ),
    )


@pytest.mark.parametrize(
    "phrase", ct.rule(CONTRACT, "fischer_adoption_claim")["blocked_phrases"]
)
def test_every_blocked_fischer_adoption_phrase_is_rejected(base, phrase):
    assert "fischer_adoption_claim" in violations_of(
        base,
        change(
            section="exp.1.bullet.1",
            before=BULLET_PROTOTYPE,
            after=f"{BULLET_PROTOTYPE}, {phrase}",
        ),
    )


def test_the_same_adoption_wording_is_allowed_on_the_growth_directorate(base):
    """`allowed_for: growth_directorate` — there the systems are in daily use."""
    assert "fischer_adoption_claim" not in violations_of(
        base,
        change(
            section="exp.0.bullet.2",
            before=BULLET_DASHBOARDS,
            after=BULLET_DASHBOARDS + " in production",
        ),
    )


def test_make_may_not_be_attributed_to_fischer(base):
    assert "attribute_make_to_fischer" in violations_of(
        base,
        change(
            section="exp.1.bullet.1",
            before=BULLET_PROTOTYPE,
            after=BULLET_PROTOTYPE + " with Make",
        ),
    )


def test_make_in_an_unattributed_skills_line_is_fine(base):
    assert "attribute_make_to_fischer" not in violations_of(
        base,
        change(
            op=cs.ADD_TERM_TO_EXISTING_LINE,
            section="skills.1",
            before=SKILL_TOOLS,
            after=SKILL_TOOLS + ", make scenarios",
            source=cs.BASE,
            source_line=SKILL_TOOLS,
        ),
    )


def test_vertex_may_not_be_put_on_fischer(base):
    assert "vertex_on_fischer" in violations_of(
        base,
        change(
            section="exp.1.bullet.0",
            before=ANCHORS["fischer_ai_solutions"]["en"],
            after=ANCHORS["fischer_ai_solutions"]["en"] + " on Vertex AI",
        ),
    )


def test_vertex_stays_allowed_in_the_cloud_skills_line(base):
    assert "vertex_on_fischer" not in violations_of(
        base,
        change(
            op=cs.ADD_TERM_TO_EXISTING_LINE,
            section="skills.2",
            before=SKILL_CLOUD,
            after=SKILL_CLOUD + ", vertex pipelines",
            source=cs.BASE,
            source_line=SKILL_CLOUD,
        ),
    )


@pytest.mark.parametrize("spelling", ["AI-assisted", "AI assisted", "ai_assisted"])
def test_ai_assisted_is_never_written(base, spelling):
    assert "write_ai_assisted" in violations_of(
        base,
        change(
            section="exp.0.bullet.2",
            before=BULLET_DASHBOARDS,
            after=f"{spelling} dashboards for the weekly review",
        ),
    )


@pytest.mark.parametrize(
    "address", ["identity.1", "exp.0.header", "exp.0.subtitle", "structural.0", "heading.2"]
)
def test_the_identity_block_and_everything_structural_is_untouchable(base, address):
    line = base.line(address)
    assert "touch_identity_block" in violations_of(
        base, change(section=address, before=line.text, after=line.text + " (tailored)")
    )


@pytest.mark.parametrize(
    "escalation", CONTRACT["evidence"]["scope_unchanged"]["blocked_escalations"]
)
def test_no_swap_may_enlarge_the_claim(base, escalation):
    weaker, stronger = [part.strip() for part in escalation.split("->")]
    assert "scope_unchanged" in violations_of(
        base,
        change(
            section="exp.0.bullet.2",
            before=f"{weaker} the weekly reporting flow",
            after=f"{stronger} the weekly reporting flow",
        ),
    )


def test_an_operation_outside_the_allowed_list_is_not_a_change(base):
    assert "allowed" in violations_of(base, change(op="rewrite_document"))


def test_an_anchor_is_never_a_secondary_bullet(base):
    assert "allowed" in violations_of(
        base, change(op=cs.DROP_SECONDARY_BULLET, section="exp.1.bullet.0", after="")
    )


def test_a_reorder_must_be_a_permutation_of_the_same_lines(base):
    assert "allowed" in violations_of(
        base,
        change(
            op=cs.REORDER_SKILL_CATEGORIES,
            section="skills",
            before="skills.0, skills.1",
            after="skills.0",
        ),
    )


def test_enforce_raises_and_names_the_rule(base):
    with pytest.raises(ct.ContractError) as caught:
        ct.enforce(
            base,
            cs.ChangeSet(changes=(change(section="summary.0", after="rewritten"),)),
            contract=CONTRACT,
        )
    assert "edit_summary" in str(caught.value)
    assert any(v.rule == "edit_summary" for v in caught.value.violations)


def test_a_forbidden_rule_with_no_checker_fails_the_run(base):
    """Adding a rule to the YAML and forgetting the code must be loud."""
    invented = dict(CONTRACT)
    invented["forbidden"] = list(CONTRACT["forbidden"]) + [{"id": "no_emoji", "rule": "no emoji"}]
    rules = {v.rule for v in ct.check(base, cs.ChangeSet(), contract=invented)}
    assert "no_emoji" in rules


@pytest.mark.parametrize(
    ("case", "rule_id"),
    [
        (
            {
                "section": "exp.1.bullet.1",
                "before": "Built a prototype in production for one client",
                "after": "Built a retrieval prototype in production for one client",
            },
            "fischer_adoption_claim",
        ),
        (
            {
                "section": "exp.1.bullet.1",
                "before": BULLET_PROTOTYPE + " with Make",
                "after": BULLET_PROTOTYPE + " with Make",
            },
            "attribute_make_to_fischer",
        ),
        (
            {
                "section": "exp.1.bullet.1",
                "before": BULLET_PROTOTYPE + " on Vertex AI",
                "after": BULLET_PROTOTYPE + " on Vertex AI",
            },
            "vertex_on_fischer",
        ),
        (
            {
                "section": "exp.0.bullet.2",
                "before": "Led the weekly review dashboards",
                "after": "Led the weekly review dashboards end to end",
            },
            "scope_unchanged",
        ),
    ],
)
def test_a_change_may_not_invent_the_before_it_is_judged_against(base, case, rule_id):
    """Four rules are diffs of `before` against `after`, and nothing checked `before`.

    A change that writes its own `before` chooses what those four rules compare
    against — declare the Fischer line already said "in production" and the
    production claim reads as pre-existing — while the projection applies
    `after` regardless. `before` is evidence, so it is the base's line or the
    change is unsourced.
    """
    caught = violations_of(base, change(**case))
    assert "claim_without_evidence" in caught, f"a false `before` also defeated {rule_id}"
    assert base.line(case["section"]).text != case["before"]


def test_a_dropped_bullet_may_not_carry_replacement_text(base, tmp_path):
    """The projection read it as a removal, the renderer wrote the text anyway."""
    sneak = change(
        op=cs.DROP_SECONDARY_BULLET,
        section="exp.1.bullet.1",
        before=BULLET_PROTOTYPE,
        after="Built a retrieval prototype, in production, adopted by the client",
    )
    assert "allowed" in violations_of(base, sneak)

    out = tmp_path / "out" / "cv.docx"
    render.write(base, cs.ChangeSet(changes=(sneak,)), out)
    again = bases.load(out)
    assert again.line("exp.1.bullet.1").text == BULLET_PROTOTYPE
    assert cs.project(base, cs.ChangeSet(changes=(sneak,))).of("exp.1.bullet.1") == BULLET_PROTOTYPE


@pytest.mark.parametrize(
    ("section", "before", "after"),
    [
        ("skills", "identity.0, identity.1", "identity.1, identity.0"),
        ("exp.0", "exp.0.header, exp.1.header", "exp.1.header, exp.0.header"),
        ("skills", "skills.0, structural.0", "structural.0, skills.0"),
        ("exp.0", "exp.0.bullet.0, exp.1.bullet.0", "exp.1.bullet.0, exp.0.bullet.0"),
    ],
)
def test_a_reorder_may_not_move_anything_but_its_own_skills_and_bullets(
    base, section, before, after
):
    """A reorder was exempt from `touch_identity_block` and its addresses were never checked."""
    assert "touch_identity_block" in violations_of(
        base,
        change(op=cs.REORDER_LINES_WITHIN_SECTION, section=section, before=before, after=after),
    )


@pytest.mark.parametrize("control", ["\n", "\r", "\t", "\v", "\u2028"])
def test_a_control_character_in_after_is_a_new_line_and_is_rejected(base, control):
    """A newline in `after` is a real second line in the .docx that every check counts as one."""
    assert "add_new_bullet" in violations_of(
        base,
        change(
            section="exp.0.bullet.2",
            before=BULLET_DASHBOARDS,
            after=f"{BULLET_DASHBOARDS}{control}Owned the agentic orchestration platform",
        ),
    )


# --- rendering ------------------------------------------------------------


def test_rendering_edits_the_base_and_leaves_everything_else_alone(base, tmp_path):
    approved = cs.ChangeSet(
        changes=(
            change(
                section="exp.0.bullet.2",
                before=BULLET_DASHBOARDS,
                after="Built reporting dashboards for the weekly review",
                source=cs.INVENTORY,
                source_line="Reporting dashboards were rebuilt during the quarterly review.",
            ),
            change(op=cs.DROP_SECONDARY_BULLET, section="exp.1.bullet.1", after=""),
        )
    )
    out = tmp_path / "out" / "cv.docx"
    written = render.write(base, approved, out)

    assert written.path.exists()
    assert (written.changed, written.removed) == (1, 1)

    again = bases.load(out)
    assert again.summary_text == SUMMARY_TEXT
    assert again.line("identity.0").text == "Test Person"
    assert again.line("exp.0.bullet.2").text == "Built reporting dashboards for the weekly review"
    assert [line.address for line in again.lines if line.group == "exp.1"] == [
        "exp.1.header",
        "exp.1.subtitle",
        "exp.1.bullet.0",
    ]


def test_rendering_keeps_the_bold_lead_in_of_a_skills_line(base, tmp_path):
    out = tmp_path / "out" / "cv.docx"
    render.write(
        base,
        cs.ChangeSet(
            changes=(
                change(
                    op=cs.ADD_TERM_TO_EXISTING_LINE,
                    section="skills.0",
                    before=SKILL_DATA,
                    after=SKILL_DATA + ", dbt",
                    source=cs.INVENTORY,
                    source_line="dbt",
                ),
            )
        ),
        out,
    )
    reopened = Document(str(out))
    skills = [p for p in reopened.paragraphs if p.text.startswith("Data:")][0]
    assert skills.text == SKILL_DATA + ", dbt"
    # `any(run.bold ...)` held even under total flattening, because a flattened
    # paragraph is one bold run. What the localised edit actually promises is
    # that the bold stayed on the category and did not spread over the items.
    assert "".join(run.text for run in skills.runs if run.bold) == "Data: "


def test_a_change_that_says_nothing_does_not_flatten_the_paragraph(base, tmp_path):
    """`before == after` used to collapse every run into run 0, bold and all."""
    out = tmp_path / "out" / "cv.docx"
    written = render.write(
        base,
        cs.ChangeSet(
            changes=(change(section="skills.0", before=SKILL_DATA, after=SKILL_DATA),)
        ),
        out,
    )
    assert written.changed == 0, "a change that changes nothing is not applied"

    skills = [p for p in Document(str(out)).paragraphs if p.text.startswith("Data:")][0]
    assert skills.text == SKILL_DATA
    assert "".join(run.text for run in skills.runs if run.bold) == "Data: "


def test_rendering_refuses_a_paragraph_that_does_not_say_what_before_says(base):
    """The base moved under the changeset. Rewriting run 0 wrote a line nobody checked."""
    with pytest.raises(render.BaseMismatch):
        render.apply(
            base,
            cs.ChangeSet(
                changes=(
                    change(
                        section="skills.0",
                        before="Data: something the base never said",
                        after="Data: sql, python, dbt",
                    ),
                )
            ),
        )
    assert base.document.paragraphs[base.line("skills.0").index].text == SKILL_DATA


def test_the_reorder_moves_the_paragraphs_not_only_the_text(base, tmp_path):
    out = tmp_path / "out" / "cv.docx"
    render.write(
        base,
        cs.ChangeSet(
            changes=(
                change(
                    op=cs.REORDER_SKILL_CATEGORIES,
                    section="skills",
                    before="skills.0, skills.1, skills.2",
                    after="skills.2, skills.0, skills.1",
                ),
            )
        ),
        out,
    )
    again = bases.load(out)
    assert [line.text for line in again.skills] == [SKILL_CLOUD, SKILL_DATA, SKILL_TOOLS]


def test_the_filename_never_names_the_role(tmp_path):
    path = render.output_path(
        CONTRACT,
        company="Acme",
        title="Senior Data Analyst",
        fingerprint="abc123",
        root=tmp_path,
    )
    assert path.name == CONTRACT["review"]["output"]["filename"]
    assert "analyst" not in path.name.lower()
    assert path.parent.name == "Acme - Senior Data Analyst"


def test_the_output_folder_falls_back_to_the_fingerprint(tmp_path):
    path = render.output_path(CONTRACT, fingerprint="deadbeefcafe0000", root=tmp_path)
    assert path.parent.name == "deadbeefcafe"


@pytest.mark.parametrize("company", [".", "..", "...", " . ", ".hidden"])
def test_a_scraped_company_name_cannot_walk_out_of_the_output_folder(tmp_path, company):
    """`company` and `title` are scraped, and a folder of "." is not a new folder.

    "." wrote the CV straight into the folder that holds the six bases and the
    experience inventory; ".." resolved a level above it.
    """
    root = tmp_path / "קורות חיים"
    path = render.output_path(CONTRACT, company=company, fingerprint="abc123def456", root=root)

    assert path.parent.name not in (".", "..", "")
    assert not path.parent.name.startswith(".")
    assert path.resolve().is_relative_to(root.resolve())
    assert path.parent.resolve() != root.resolve()


def test_an_output_path_that_escaped_the_root_is_an_error(tmp_path, monkeypatch):
    monkeypatch.setattr(render, "folder_name", lambda **_: "../elsewhere")
    with pytest.raises(render.UnsafeOutputPath):
        render.output_path(CONTRACT, company="Acme", root=tmp_path)


def test_writing_never_silently_overwrites_a_document_being_edited(bases_dir, tmp_path):
    """`format: docx` exists because Noam edits the file in Word afterwards.

    A second `desk tailor --write` used to destroy those edits with no prompt
    and no backup.
    """
    out = tmp_path / "out" / "cv.docx"
    approved = cs.ChangeSet(
        changes=(
            change(
                section="exp.0.bullet.2",
                before=BULLET_DASHBOARDS,
                after="Built reporting dashboards for the weekly review",
                source=cs.INVENTORY,
                source_line="Reporting dashboards were rebuilt during the quarterly review.",
            ),
        )
    )
    render.write(bases.load_for("ai_builder", directory=bases_dir), approved, out)
    out.write_bytes(b"an evening of Noam's edits, in Word")

    with pytest.raises(render.OutputExists):
        render.write(bases.load_for("ai_builder", directory=bases_dir), approved, out)
    assert out.read_bytes() == b"an evening of Noam's edits, in Word"

    render.write(bases.load_for("ai_builder", directory=bases_dir), approved, out, force=True)
    assert bases.load(out).line("exp.0.bullet.2").text.startswith("Built reporting")
    assert [p.name for p in out.parent.iterdir()] == [out.name], "no staging file left behind"


def test_the_diff_shows_the_evidence_behind_every_change(base):
    rows = render.diff(
        base,
        cs.ChangeSet(
            changes=(
                change(
                    section="exp.1.bullet.1",
                    before=BULLET_PROTOTYPE,
                    after="Built a retrieval prototype for one public-sector client",
                    source=cs.INVENTORY,
                    source_line="Retrieval work used an open-source vector store.",
                ),
            ),
            missing_requirements=("kubernetes",),
        ),
    )
    joined = "\n".join(rows)
    assert "fischer_technologies" in joined
    assert "inventory: Retrieval work" in joined
    assert "gap  kubernetes" in joined


# --- the run --------------------------------------------------------------


@dataclass
class FakeClient:
    """A model that answers from a script. No key, no network, no cassette."""

    payloads: dict[str, Any]
    name: str = "fake"
    stages: list[str] = field(default_factory=list)

    def complete(self, req: Any, route: Any) -> LLMResponse:
        self.stages.append(req.stage)
        text = json.dumps(self.payloads[req.stage], ensure_ascii=False)
        return LLMResponse(
            text=text,
            usage=Usage(input_tokens=10, output_tokens=10),
            model=route.model,
            stage=req.stage,
            parsed=json.loads(text),
        )


def analysis_for(*requirements: str) -> Analysis:
    return Analysis(
        fingerprint="fp-1",
        title="AI Builder",
        company="Acme",
        family=Family(family="ai_builder", confidence=0.9, reason="agentic work"),
        requirements=tuple(
            Requirement(text=r, kind="skill", evidence=f"the posting says {r}")
            for r in requirements
        ),
        fit=Fit(score=0.8, rationale="close", channel="person"),
    )


def run_tailor(ctx, bases_dir, inventory, payloads, **kwargs):
    ctx.gateway.client = FakeClient(payloads)
    return tl.tailor(
        kwargs.pop("analysis", analysis_for("reporting dashboards")),
        ctx=ctx,
        bases_dir=bases_dir,
        inventory_path=inventory,
        contract=CONTRACT,
        **kwargs,
    )


def _clean_payload() -> dict[str, Any]:
    return {
        "tailor_cv": {
            "changes": [
                {
                    "op": cs.SWAP_TERMINOLOGY,
                    "section": "exp.0.bullet.2",
                    "before": BULLET_DASHBOARDS,
                    "after": "Built reporting dashboards for the weekly review",
                    "source": "inventory",
                    "source_line": "Reporting dashboards were rebuilt during the quarterly review.",
                }
            ],
            "missing_requirements": [],
        },
        "verify_no_fabrication": {"ok": True, "unsupported": []},
    }


def test_a_run_proposes_enforces_verifies_and_writes_nothing(ctx, bases_dir, inventory):
    result = run_tailor(ctx, bases_dir, inventory, _clean_payload())

    assert ctx.gateway.client.stages == ["tailor_cv", "verify_no_fabrication"]
    assert result.ok and result.path is None
    assert result.base.sha256 == bases.sha256_of(result.base.path)
    assert result.projected.of("summary.0") == SUMMARY_TEXT
    assert result.projected.of("exp.0.bullet.2").startswith("Built reporting")


def test_a_requirement_the_inventory_does_not_cover_never_enters_the_cv(
    ctx, bases_dir, inventory
):
    """`missing_requirement.not_in_inventory: report`, and `never: fabricate`."""
    result = run_tailor(
        ctx,
        bases_dir,
        inventory,
        _clean_payload(),
        analysis=analysis_for("reporting dashboards", "kubernetes on bare metal"),
    )
    assert not any("kubernetes" in t.lower() for t in result.projected.texts())
    assert any("kubernetes" in gap.lower() for gap in result.gaps)
    assert any("not in the inventory" in gap for gap in result.gaps)


def test_a_requirement_the_inventory_covers_is_reported_as_the_other_kind_of_gap(
    ctx, bases_dir, inventory
):
    result = run_tailor(
        ctx,
        bases_dir,
        inventory,
        _clean_payload(),
        analysis=analysis_for("open-source vector store"),
    )
    assert any("in the inventory, not in the CV" in gap for gap in result.gaps)


def test_a_model_that_breaks_the_contract_fails_the_run_before_anything_is_written(
    ctx, bases_dir, inventory
):
    payload = _clean_payload()
    payload["tailor_cv"]["changes"] = [
        {
            "op": cs.SWAP_TERMINOLOGY,
            "section": "summary.0",
            "before": SUMMARY_TEXT,
            "after": "A summary tuned to this one posting.",
            "source": "base",
            "source_line": SUMMARY_TEXT,
        }
    ]
    with pytest.raises(ct.ContractError) as caught:
        run_tailor(ctx, bases_dir, inventory, payload)

    assert any(v.rule == "edit_summary" for v in caught.value.violations)
    assert ctx.gateway.client.stages == ["tailor_cv"], "the verifier is never reached"


def test_a_changeset_the_verifier_refuses_never_becomes_a_document(ctx, bases_dir, inventory):
    payload = _clean_payload()
    payload["verify_no_fabrication"] = {"ok": False, "unsupported": ["exp.0.bullet.2"]}
    with pytest.raises(tl.Fabrication) as caught:
        run_tailor(ctx, bases_dir, inventory, payload)
    assert caught.value.unsupported == ("exp.0.bullet.2",)


def test_a_reorder_only_changeset_skips_the_verifier(ctx, bases_dir, inventory):
    payload = _clean_payload()
    payload["tailor_cv"]["changes"] = [
        {
            "op": cs.REORDER_SKILL_CATEGORIES,
            "section": "skills",
            "before": "skills.0, skills.1, skills.2",
            "after": "skills.2, skills.0, skills.1",
            "source": "base",
            "source_line": "the skills section",
        }
    ]
    result = run_tailor(ctx, bases_dir, inventory, payload)
    assert ctx.gateway.client.stages == ["tailor_cv"]
    assert result.ok


def test_a_posting_that_matched_no_family_is_not_tailored(ctx, bases_dir, inventory):
    with pytest.raises(tl.NoFamily):
        run_tailor(ctx, bases_dir, inventory, _clean_payload(), analysis=Analysis("fp-2"))


def test_a_missing_inventory_is_reported_and_not_an_error(ctx, bases_dir, tmp_path):
    result = run_tailor(ctx, bases_dir, tmp_path / "nope.md", _clean_payload())
    assert any("inventory" in note for note in result.notes)
    assert tl.read_inventory(tmp_path / "nope.md") == ""


def test_the_stored_analysis_round_trips_into_the_frozen_type(ctx):
    analysis = analysis_for("reporting dashboards")
    ctx.store.put_analysis(
        analysis.fingerprint,
        analysis.as_json(),
        family="ai_builder",
        score=0.8,
        channel="person",
        rationale="close",
        stopped_at="",
        now="2026-08-18T00:00:00",
    )
    loaded = tl.load_analysis(ctx.store, analysis.fingerprint)
    assert loaded is not None and loaded.family.family == "ai_builder"
    assert tl.load_analysis(ctx.store, "nothing") is None


# --- the prompts ----------------------------------------------------------


def test_the_prompts_are_versioned_files_and_render(base):
    from desk import prompts

    proposal = prompts.load("tailor", "tailor_cv", 1)
    rendered = proposal.render(
        ops="swap_terminology",
        rules=tl.render_rules(CONTRACT),
        family="ai_builder",
        language="en",
        requirements="- (must, skill) reporting dashboards",
        gaps="(none)",
        lines=tl.render_lines(base),
        inventory="(none)",
    )
    assert "edit_summary" in rendered
    assert "[lock] summary.0" in rendered
    assert "[edit] skills.0" in rendered

    verifier = prompts.load("tailor", "verify_no_fabrication", 1)
    assert verifier.render(changes="x", base="y", inventory="z")
    assert proposal.sha256 != verifier.sha256


# --- the real bases, when this machine has them ---------------------------

REAL_BASES = Path(CONTRACT["inputs"]["bases_dir"]).expanduser()


@pytest.mark.skipif(not REAL_BASES.exists(), reason="the approved bases are not on this machine")
def test_the_approved_bases_parse_and_carry_their_anchors():
    """Structure only. Nothing from these files is asserted on, printed or stored."""
    catalog = bases.catalog(REAL_BASES)
    assert len(catalog) == 6

    for entry in catalog:
        base = bases.load(entry.path, family=entry.family, language=entry.language)
        assert base.summary, entry.family
        assert base.skills, entry.family
        assert base.employers, entry.family
        located = [a for a in ct.locate_anchors(base, CONTRACT) if a.located]
        assert len(located) == len(CONTRACT["anchors"]), entry.family


# --- the write goes through the registry ----------------------------------
#
# `desk tailor --write` is the only write-local act in the daily run, and it is
# performed by dispatching `write_tailored_cv` rather than by calling the
# renderer directly. These tests exist because that was not true until now: the
# registry gated a tool nothing in the daily path called, which is a guarantee
# about a code path instead of about the system.


def tailoring_contract(bases_dir: Path, out: Path) -> dict[str, Any]:
    """The real contract with both ends pointed at tmp_path."""
    contract = copy.deepcopy(CONTRACT)
    contract["inputs"]["bases_dir"] = str(bases_dir)
    contract["review"]["output"]["dir"] = f"{out}/{render.FOLDER_PLACEHOLDER}/"
    return contract


APPROVED = cs.ChangeSet(
    changes=(
        cs.Change(
            op=cs.SWAP_TERMINOLOGY,
            section="exp.0.bullet.2",
            before=BULLET_DASHBOARDS,
            after="Built reporting dashboards for the weekly review",
            source=cs.INVENTORY,
            source_line="Reporting dashboards were rebuilt during the quarterly review.",
        ),
    )
)


def dispatch_write(ctx, base, **overrides):
    args = {
        "fingerprint": "deadbeefcafe0000",
        "family": base.family,
        "language": base.language,
        "base_sha256": base.sha256,
        "changeset": APPROVED.as_json(),
        "company": "Acme",
        "title": "AI Engineer",
    }
    args.update(overrides)
    return registry.dispatch("write_tailored_cv", args, ctx)


def test_the_document_is_cut_by_dispatching_the_registered_tool(make_ctx, bases_dir, tmp_path):
    out = tmp_path / "documents"
    ctx = make_ctx(approval_token="local-run")
    ctx.contract = tailoring_contract(bases_dir, out)

    result = dispatch_write(ctx, bases.load_for("ai_builder", directory=bases_dir))

    assert result.ok, result.error
    written = Path(result.content["path"])
    assert written.exists() and written.is_relative_to(out)
    assert result.content["changed"] == 1

    again = bases.load(written)
    assert again.line("exp.0.bullet.2").text == "Built reporting dashboards for the weekly review"
    assert again.summary_text == SUMMARY_TEXT


def test_a_dry_run_carries_no_token_and_the_write_is_denied(make_ctx, bases_dir, tmp_path):
    """The `--write` flag is the approval token, and the denial is not the caller's to make."""
    out = tmp_path / "documents"
    ctx = make_ctx(approval_token=None)
    ctx.contract = tailoring_contract(bases_dir, out)

    result = dispatch_write(ctx, bases.load_for("ai_builder", directory=bases_dir))

    assert result.denied and not result.ok
    assert "approval token" in str(result.error)
    assert not out.exists(), "a denied write still put a document on disk"


def test_a_base_edited_since_the_changeset_was_cut_is_refused(make_ctx, bases_dir, tmp_path):
    out = tmp_path / "documents"
    ctx = make_ctx(approval_token="local-run")
    ctx.contract = tailoring_contract(bases_dir, out)
    base = bases.load_for("ai_builder", directory=bases_dir)

    result = dispatch_write(ctx, base, base_sha256="0" * 64)

    assert not result.ok and not result.denied
    assert "BaseMismatch" in str(result.error)
    assert not out.exists()


def test_the_dispatched_write_never_silently_replaces_a_document(make_ctx, bases_dir, tmp_path):
    out = tmp_path / "documents"
    ctx = make_ctx(approval_token="local-run")
    ctx.contract = tailoring_contract(bases_dir, out)
    base = bases.load_for("ai_builder", directory=bases_dir)

    first = dispatch_write(ctx, base)
    assert first.ok, first.error
    Path(first.content["path"]).write_bytes(b"an evening of edits in Word")

    second = dispatch_write(ctx, bases.load_for("ai_builder", directory=bases_dir))
    assert not second.ok
    assert "OutputExists" in str(second.error)
    assert Path(first.content["path"]).read_bytes() == b"an evening of edits in Word"

    forced = dispatch_write(ctx, bases.load_for("ai_builder", directory=bases_dir), force=True)
    assert forced.ok, forced.error
    assert Path(forced.content["path"]).read_bytes() != b"an evening of edits in Word"


def test_no_argument_of_the_write_tool_names_a_destination():
    """The destination is derived from the contract; a model never supplies one.

    `evals/guardrails.py` asserts the same property across the whole registry.
    It is repeated here because this is the tool that writes to disk, and the
    guardrail would still pass if this schema quietly grew a `path`.
    """
    schema = registry.get("write_tailored_cv").input_schema
    assert not set(schema["properties"]) & set(_REACH_PARAMS)
